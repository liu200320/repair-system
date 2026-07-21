"""
维修记录 Word 导出 — 单页紧凑版
每条记录严格控制在一页内（A4，1.2cm 边距）
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.models.repair import RepairRecord, PhotoPhase
from app.core.config import UPLOAD_DIR

STATUS_LABELS = {
    "pending":     "待维修",
    "in_progress": "维修中",
    "completed":   "已完成",
}
PHASE_ORDER = [
    (PhotoPhase.before, "维修前"),
    (PhotoPhase.during, "维修中"),
    (PhotoPhase.after,  "维修后"),
]
MAX_PHOTOS_PER_PHASE = 2   # 每阶段最多2张，共6张
PHOTO_WIDTH = 1.7           # 英寸


def _set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _hdr(cell, text, bg="DBEAFE"):
    cell.text = text
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _zero_spacing(p)
    if p.runs:
        p.runs[0].bold = True; p.runs[0].font.size = Pt(9)
    _set_cell_bg(cell, bg)


def _val(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = str(text) if text is not None else ""
    p = cell.paragraphs[0]; p.alignment = align
    _zero_spacing(p)
    if p.runs:
        p.runs[0].font.size = Pt(9)


def _zero_spacing(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)


def _insert_photo(cell, img_path, caption):
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _zero_spacing(p)
    try:
        p.add_run().add_picture(img_path, width=Inches(PHOTO_WIDTH))
        cap = cell.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _zero_spacing(cap)
        if cap.runs:
            cap.runs[0].font.size = Pt(7)
            cap.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    except Exception:
        p.text = "[图片加载失败]"


def _add_page_break(doc):
    p = doc.add_paragraph()
    _zero_spacing(p)
    run = p.add_run()
    br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_doc_margins(doc):
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.2)


def _write_record(doc, record: RepairRecord):
    # ── 标题
    tp = doc.add_paragraph()
    _zero_spacing(tp); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("云南农业职业技术学院安防设备维修记录")
    tr.bold = True; tr.font.size = Pt(13)

    # ── 基本信息（3行×4列）
    t = doc.add_table(rows=3, cols=4); t.style = "Table Grid"
    col_w = [Cm(2.4), Cm(6.0), Cm(2.4), Cm(5.8)]
    for row in t.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_w[ci]

    _hdr(t.rows[0].cells[0], "工单编号"); _val(t.rows[0].cells[1], record.record_no)
    _hdr(t.rows[0].cells[2], "维修日期"); _val(t.rows[0].cells[3], record.repair_date, WD_ALIGN_PARAGRAPH.CENTER)
    _hdr(t.rows[1].cells[0], "维修点位"); _val(t.rows[1].cells[1], record.location)
    _hdr(t.rows[1].cells[2], "维修人员"); _val(t.rows[1].cells[3], record.repairer or "")
    _hdr(t.rows[2].cells[0], "当前状态"); _val(t.rows[2].cells[1], STATUS_LABELS.get(record.status.value, record.status.value), WD_ALIGN_PARAGRAPH.CENTER)
    _hdr(t.rows[2].cells[2], "故障描述"); _val(t.rows[2].cells[3], (record.description or "")[:200])

    # ── 维修内容
    t2 = doc.add_table(rows=1, cols=2); t2.style = "Table Grid"
    t2.rows[0].cells[0].width = Cm(2.4); t2.rows[0].cells[1].width = Cm(14.2)
    _hdr(t2.rows[0].cells[0], "维修内容")
    _val(t2.rows[0].cells[1], (record.repair_content or "")[:300])

    # ── 照片（三阶段×最多2张，3列布局）
    tp2 = doc.add_paragraph(); _zero_spacing(tp2)
    r2 = tp2.add_run("维修照片"); r2.bold = True; r2.font.size = Pt(10)

    # 每阶段最多2张，共6格，用3列（每列放一个阶段）
    pt = doc.add_table(rows=2, cols=3); pt.style = "Table Grid"
    for row in pt.rows:
        for cell in row.cells:
            cell.width = Cm(5.6)

    for col_i, (phase, label) in enumerate(PHASE_ORDER):
        phase_photos = [p for p in record.photos if p.phase == phase][:MAX_PHOTOS_PER_PHASE]
        # 行0：阶段标题
        _hdr(pt.rows[0].cells[col_i], label)
        # 行1：图片（最多2张横排，子表格）
        cell = pt.rows[1].cells[col_i]
        if phase_photos:
            sub = cell.add_table(rows=1, cols=len(phase_photos))
            sub.style = "Table Grid"
            for pi, photo in enumerate(phase_photos):
                img_path = os.path.join(UPLOAD_DIR, photo.filename)
                sc = sub.rows[0].cells[pi]
                sc.width = Cm(5.6 / len(phase_photos))
                if os.path.exists(img_path):
                    _insert_photo(sc, img_path, photo.original_name or f"图{pi+1}")
                else:
                    sc.paragraphs[0].text = "（文件不存在）"
        else:
            p = cell.paragraphs[0]; p.text = "无图片"; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _zero_spacing(p)

    # ── 签字栏
    sp = doc.add_paragraph(f"维修人员：___________    日期：{record.repair_date}")
    sp.alignment = WD_ALIGN_PARAGRAPH.RIGHT; _zero_spacing(sp)
    if sp.runs: sp.runs[0].font.size = Pt(9)


def export_to_word(record: RepairRecord, output_path: str) -> str:
    doc = Document(); _set_doc_margins(doc)
    _write_record(doc, record)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def export_range_to_word(records, start_date, end_date, output_path):
    doc = Document(); _set_doc_margins(doc)
    for i, record in enumerate(records):
        _write_record(doc, record)
        if i < len(records) - 1:
            _add_page_break(doc)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
