"""
网络基础设施日常巡检登记表 Word 导出 — 单页紧凑版
"""
import os
from math import ceil
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.models.network_inspection import NetworkInspectionRecord
from app.core.config import UPLOAD_DIR

MAX_PHOTOS = 6
PHOTO_WIDTH = 1.7


def _set_cell_bg(cell, fill_hex):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _zero(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)


def _hdr(cell, text, bg="D1E8FF", align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = text
    p = cell.paragraphs[0]; p.alignment = align; _zero(p)
    if p.runs:
        p.runs[0].bold = True; p.runs[0].font.size = Pt(9)
    _set_cell_bg(cell, bg)


def _val(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = str(text) if text is not None else ""
    p = cell.paragraphs[0]; p.alignment = align; _zero(p)
    if p.runs: p.runs[0].font.size = Pt(9)


def _insert_photo(cell, img_path, caption):
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; _zero(p)
    try:
        p.add_run().add_picture(img_path, width=Inches(PHOTO_WIDTH))
        cap = cell.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; _zero(cap)
        if cap.runs:
            cap.runs[0].font.size = Pt(7)
            cap.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    except Exception:
        p.text = "[图片加载失败]"


def _set_doc_margins(doc):
    s = doc.sections[0]
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.2)


def _write_single_record(doc, record: NetworkInspectionRecord):
    # 标题
    tp = doc.add_paragraph(); _zero(tp); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("云南农业职业技术学院校园网络基础设施日常巡检登记表")
    tr.bold = True; tr.font.size = Pt(13)

    # 基本信息（2行×4列）
    info = doc.add_table(rows=2, cols=4); info.style = "Table Grid"
    for row in info.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = [Cm(2.4), Cm(6.0), Cm(2.4), Cm(5.8)][ci]
    _hdr(info.rows[0].cells[0], "巡检单号"); _val(info.rows[0].cells[1], record.record_no)
    _hdr(info.rows[0].cells[2], "巡检日期"); _val(info.rows[0].cells[3], record.inspect_date, WD_ALIGN_PARAGRAPH.CENTER)
    _hdr(info.rows[1].cells[0], "巡检地点"); _val(info.rows[1].cells[1], record.location)
    _hdr(info.rows[1].cells[2], "巡检人员"); _val(info.rows[1].cells[3], record.inspector or "")

    # 状态表（2行×4列）
    st = doc.add_table(rows=2, cols=4); st.style = "Table Grid"
    for row in st.rows:
        for cell in row.cells: cell.width = Cm(4.15)
    _hdr(st.rows[0].cells[0], "网络线路情况"); _hdr(st.rows[0].cells[1], "接入设备情况")
    _hdr(st.rows[0].cells[2], "终端网络情况"); _hdr(st.rows[0].cells[3], "其他设备")
    _val(st.rows[1].cells[0], record.line_status or "")
    _val(st.rows[1].cells[1], record.device_status or "")
    _val(st.rows[1].cells[2], record.terminal_status or "")
    _val(st.rows[1].cells[3], record.other_device or "")

    # 故障描述 + 维修内容（合并为一个2行×2列表格）
    dt = doc.add_table(rows=2, cols=2); dt.style = "Table Grid"
    for row in dt.rows:
        dt.rows[0].cells[0].width = Cm(2.4); dt.rows[0].cells[1].width = Cm(14.2)
    _hdr(dt.rows[0].cells[0], "故障描述"); _val(dt.rows[0].cells[1], (record.fault_description or "")[:300])
    _hdr(dt.rows[1].cells[0], "维修内容"); _val(dt.rows[1].cells[1], (record.repair_content or "")[:300])

    # 照片（最多6张，3列）
    photos = list(record.photos)[:MAX_PHOTOS]
    extra = len(record.photos) - MAX_PHOTOS if len(record.photos) > MAX_PHOTOS else 0
    if photos:
        lp = doc.add_paragraph(); _zero(lp)
        lr = lp.add_run(f"现场照片（共{len(record.photos)}张{f'，仅显示前{MAX_PHOTOS}张' if extra > 0 else ''}）")
        lr.bold = True; lr.font.size = Pt(10)

        n_rows = ceil(len(photos) / 3)
        pt = doc.add_table(rows=n_rows, cols=3); pt.style = "Table Grid"
        for row in pt.rows:
            for cell in row.cells: cell.width = Cm(5.6)
        for pr in range(n_rows):
            for pc in range(3):
                idx = pr * 3 + pc; cell = pt.rows[pr].cells[pc]
                if idx < len(photos):
                    ph = photos[idx]
                    img_path = os.path.join(UPLOAD_DIR, ph.filename)
                    caption = f"图{ph.photo_index} {ph.original_name or ''}"
                    if os.path.exists(img_path):
                        _insert_photo(cell, img_path, caption)
                    else:
                        cell.paragraphs[0].text = f"（图{ph.photo_index} 不存在）"
                else:
                    _zero(cell.paragraphs[0])


def export_network_inspection_to_word(record: NetworkInspectionRecord, output_path: str) -> str:
    doc = Document(); _set_doc_margins(doc)
    _write_single_record(doc, record)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def export_network_inspection_range_to_word(records, start_date, end_date, output_path):
    doc = Document(); _set_doc_margins(doc)
    for idx, record in enumerate(records):
        _write_single_record(doc, record)
        if idx < len(records) - 1:
            p = doc.add_paragraph(); _zero(p)
            run = p.add_run()
            br = OxmlElement("w:br"); br.set(qn("w:type"), "page")
            run._r.append(br)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
