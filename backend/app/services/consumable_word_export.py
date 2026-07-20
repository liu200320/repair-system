"""
耗材使用情况表 Word 导出
格式：云南农业职业技术学院电子设备网络安防耗材使用情况表
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.models.consumable import ConsumableRecord
from app.core.config import UPLOAD_DIR


# ── 工具函数（与 word_export.py 同风格）──────────────────────────

def _set_cell_bg(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _bold_cell(cell, text: str, bg: str = "DBEAFE", align=WD_ALIGN_PARAGRAPH.CENTER, font_size: int = 11):
    cell.text = text
    p = cell.paragraphs[0]
    p.alignment = align
    if p.runs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(font_size)
    _set_cell_bg(cell, bg)


def _set_cell_text(cell, text: str, align=WD_ALIGN_PARAGRAPH.LEFT, font_size: int = 11):
    cell.text = str(text) if text is not None else ""
    p = cell.paragraphs[0]
    p.alignment = align
    if p.runs:
        p.runs[0].font.size = Pt(font_size)


def _merge_row_cells(row, start_col: int, end_col: int):
    """合并同一行从 start_col 到 end_col（含）的单元格，返回合并后的单元格"""
    cell = row.cells[start_col]
    for i in range(start_col + 1, end_col + 1):
        cell = cell.merge(row.cells[i])
    return cell


def _insert_photo(cell, img_path: str, caption: str, width_inches: float = 2.8):
    """向单元格插入图片及说明文字"""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width_inches))
        cap = cell.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    except Exception:
        p.text = "[图片加载失败]"


def _set_doc_margins(doc: Document):
    section = doc.sections[0]
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


# ── 核心导出函数 ──────────────────────────────────────────────────

def export_consumable_to_word(record: ConsumableRecord, output_path: str) -> str:
    """
    将一条耗材使用记录导出为 Word，格式严格对应原始文档：
    标题 → 地点行 → 日期行 → 表格（表头 + 明细行 + 照片区 + 备注行）
    """
    doc = Document()
    _set_doc_margins(doc)

    # ── 文档标题 ────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("云南农业职业技术学院电子设备网络安防耗材使用情况表")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # ── 耗材使用地点 + 日期 ──────────────────────────────────────
    loc_p = doc.add_paragraph()
    loc_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    loc_run = loc_p.add_run(f"耗材使用地点：{record.location}")
    loc_run.font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    date_run = date_p.add_run(f"日期：{record.use_date}")
    date_run.font.size = Pt(12)

    # ── 主表格 ──────────────────────────────────────────────────
    # 行数 = 1(表头) + len(items，最少1行) + photo_rows + 1(备注)
    items = list(record.items)
    if not items:
        # 确保至少有一行空白明细，保持表格完整性
        items = []
        n_item_rows = 1
    else:
        n_item_rows = len(items)

    photos = list(record.photos)   # 按 photo_index 排序
    # 照片每行放 2 张（合并列 0-1 / 2-3），计算需要几行
    n_photo_rows = max(1, -(-len(photos) // 2)) if photos else 1  # ceil(len/2)，至少1行

    total_rows = 1 + n_item_rows + n_photo_rows + 1  # 表头 + 明细 + 照片 + 备注
    table = doc.add_table(rows=total_rows, cols=4)
    table.style = "Table Grid"

    # 设置列宽（A4 正文宽约 15cm，分配：45% 30% 10% 15%）
    col_widths = [Cm(6.75), Cm(2.0), Cm(1.5), Cm(4.75)]
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            cell.width = col_widths[ci]

    # 1. 表头行
    header_row = table.rows[0]
    _bold_cell(header_row.cells[0], "耗材名称", bg="DBEAFE")
    _bold_cell(header_row.cells[1], "单位",     bg="DBEAFE")
    _bold_cell(header_row.cells[2], "数量",     bg="DBEAFE")
    _bold_cell(header_row.cells[3], "使用人签字", bg="DBEAFE")

    # 2. 明细行
    for i in range(n_item_rows):
        row = table.rows[1 + i]
        if i < len(items):
            item = items[i]
            _set_cell_text(row.cells[0], item.name)
            _set_cell_text(row.cells[1], item.unit or "",     align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[2], item.quantity or "", align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_text(row.cells[3], item.signer or "")
        else:
            for cell in row.cells:
                cell.text = ""

    # 3. 照片行（每行 2 张，列 0-1 合并为左图，列 2-3 合并为右图）
    photo_start = 1 + n_item_rows
    for pr in range(n_photo_rows):
        row = table.rows[photo_start + pr]
        # 合并左半（col 0-1）
        left_cell  = row.cells[0].merge(row.cells[1])
        # 合并右半（col 2-3）
        right_cell = row.cells[2].merge(row.cells[3])

        left_cell.height  = Cm(6)
        right_cell.height = Cm(6)

        left_photo_idx  = pr * 2        # 0, 2, 4 …
        right_photo_idx = pr * 2 + 1   # 1, 3, 5 …

        if left_photo_idx < len(photos):
            ph = photos[left_photo_idx]
            img_path = os.path.join(UPLOAD_DIR, ph.filename)
            caption  = f"图{ph.photo_index}  {ph.original_name or ''}"
            if os.path.exists(img_path):
                _insert_photo(left_cell, img_path, caption, width_inches=2.8)
            else:
                left_cell.paragraphs[0].text = f"（图{ph.photo_index} 文件不存在）"
        else:
            left_cell.paragraphs[0].text = ""
            left_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        if right_photo_idx < len(photos):
            ph = photos[right_photo_idx]
            img_path = os.path.join(UPLOAD_DIR, ph.filename)
            caption  = f"图{ph.photo_index}  {ph.original_name or ''}"
            if os.path.exists(img_path):
                _insert_photo(right_cell, img_path, caption, width_inches=2.8)
            else:
                right_cell.paragraphs[0].text = f"（图{ph.photo_index} 文件不存在）"
        else:
            right_cell.paragraphs[0].text = ""
            right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4. 备注行（四列全合并）
    notes_row = table.rows[photo_start + n_photo_rows]
    notes_cell = notes_row.cells[0]
    for ci in range(1, 4):
        notes_cell = notes_cell.merge(notes_row.cells[ci])

    notes_p = notes_cell.paragraphs[0]
    notes_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label_run = notes_p.add_run("备注：")
    label_run.bold = True
    label_run.font.size = Pt(11)
    content_run = notes_p.add_run(record.notes or "")
    content_run.font.size = Pt(11)

    # ── 保存 ────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
